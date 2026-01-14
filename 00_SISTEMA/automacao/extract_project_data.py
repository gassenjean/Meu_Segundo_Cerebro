import pandas as pd
from docx import Document
import os
import sys

# Paths
base_path = r"c:\Users\gasse\OneDrive\Meu_Segundo_Cerebro\02_PROJETOS\KabaK\Outlet_Expansion\recursos"
excel_path = os.path.join(base_path, "Analise de negócio Jean x Samson.xlsx")
word_path = os.path.join(base_path, "Arquérito Ativa e Kabak.docx")

def extract_excel_data(filepath):
    print(f"\n--- EXCEL EXTRACTION: {filepath} ---")
    try:
        # Load all sheets to see what we have
        xls = pd.ExcelFile(filepath)
        print(f"Sheets found: {xls.sheet_names}")
        
        # Read the first sheet assuming it has the summary
        df = pd.read_excel(xls, sheet_name=0)
        
        # Print first few rows to understand structure
        print("First 10 rows of data:")
        print(df.head(10).to_string())
        
        # Try to find key metrics if they exist in recognized columns or cells
        # This is a generic dump to the log so the AI can parse it
        print("\n--- Full Data Dump (First Sheet) ---")
        print(df.to_string())
        
    except Exception as e:
        print(f"Error reading Excel: {e}")

def extract_word_data(filepath):
    print(f"\n--- WORD EXTRACTION: {filepath} ---")
    try:
        doc = Document(filepath)
        print(f"Document paragraphs: {len(doc.paragraphs)}")
        
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        print("\n--- Document Content ---")
        print("\n".join(full_text))
        
    except Exception as e:
        print(f"Error reading Word: {e}")

if __name__ == "__main__":
    output_file = os.path.join(base_path, "final_data_extraction.txt")
    
    # Redirect stdout to file
    with open(output_file, 'w', encoding='utf-8') as f:
        sys.stdout = f
        
        if os.path.exists(excel_path):
            extract_excel_data(excel_path)
        else:
            print(f"Excel file not found at: {excel_path}")
            # Try finding it with wildcard if encoding is an issue
            files = os.listdir(base_path)
            for file in files:
                if file.endswith(".xlsx"):
                    print(f"Found alternative Excel: {file}")
                    extract_excel_data(os.path.join(base_path, file))

        if os.path.exists(word_path):
            extract_word_data(word_path)
        else:
            print(f"Word file not found at: {word_path}")
             # Try finding it with wildcard
            files = os.listdir(base_path)
            for file in files:
                if file.endswith(".docx"):
                    print(f"Found alternative Word: {file}")
                    extract_word_data(os.path.join(base_path, file))
        
        # Restore stdout
        sys.stdout = sys.__stdout__
        print(f"Extraction complete. Data written to {output_file}")
